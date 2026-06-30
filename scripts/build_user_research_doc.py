import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT = Path(r"C:\Users\ashut\OneDrive\Documents\Gym Buddy")
DATA_PATH = PROJECT / "tmp" / "research_analysis" / "gymbuddy_research_analysis.json"
OUT_DIR = PROJECT / "outputs" / "gymbuddy_user_research"
DOCX_PATH = OUT_DIR / "GymBuddy_User_Research_Report.docx"
CHART_DIR = OUT_DIR / "doc_charts"

GREEN = "1F7A4D"
DARK = "162019"
MUTED = "667168"
BLUE = "24415B"
PALE_GREEN = "F0F7EF"
PALE_BLUE = "EEF6FF"
PALE_ORANGE = "FFF8EF"
BORDER = "D9E3DB"
WHITE = "FFFFFF"


def sanitize(text):
    text = "" if text is None else str(text)
    text = text.replace("’", "'").replace("“", '"').replace("”", '"')
    text = text.replace("–", "-").replace("—", "-")
    return text.encode("ascii", "ignore").decode("ascii").strip()


def pct(value):
    return f"{round(value * 100)}%"


def font(size=28, bold=False):
    candidates = [
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, max_width, fnt):
    words = sanitize(text).split()
    lines = []
    line = ""
    for word in words:
        test = word if not line else f"{line} {word}"
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def make_bar_chart(path, title, rows, percent_key="percent", width=1200, height=680):
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34, True)
    label_font = font(22, False)
    small_font = font(20, False)
    count_font = font(22, True)

    draw.text((45, 30), sanitize(title), fill=f"#{DARK}", font=title_font)
    top = 100
    left_label = 45
    left_bar = 410
    right = width - 70
    bar_width_max = right - left_bar - 70
    bar_h = 38
    gap = 28
    max_count = max([r["count"] for r in rows] + [1])

    for idx, row in enumerate(rows):
        y = top + idx * (bar_h + gap)
        label_lines = wrap_text(draw, row["label"], 320, label_font)[:2]
        for line_idx, line in enumerate(label_lines):
            draw.text((left_label, y + line_idx * 24), line, fill=f"#{DARK}", font=label_font)

        bar_len = int((row["count"] / max_count) * bar_width_max)
        draw.rounded_rectangle(
            (left_bar, y, left_bar + bar_len, y + bar_h),
            radius=8,
            fill=f"#{GREEN}",
        )
        value = f"{row['count']} ({pct(row.get(percent_key, row.get('percent', 0)))})"
        draw.text((left_bar + bar_len + 12, y + 6), value, fill=f"#{BLUE}", font=count_font)

    draw.text((45, height - 38), "Source: GymBuddy User Research Google Form, n=20", fill=f"#{MUTED}", font=small_font)
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, header_fill=GREEN):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = sanitize(header)
        set_cell_shading(hdr[idx], header_fill)
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = sanitize(value)
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(22, 32, 25)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(sanitize(text), level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(sanitize(title))
    r.bold = True
    r.font.color.rgb = RGBColor(31, 122, 77)
    r.font.size = Pt(11)
    p.add_run("\n" + sanitize(body))
    for run in p.runs:
        run.font.name = "Calibri"
    doc.add_paragraph()


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("GymBuddy User Research")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 113, 104)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))

    charts = {
        "pain": CHART_DIR / "pain_points.png",
        "ai": CHART_DIR / "ai_trust.png",
        "features": CHART_DIR / "desired_features.png",
        "barriers": CHART_DIR / "consistency_barriers.png",
    }
    make_bar_chart(charts["pain"], "Biggest In-Gym Pain Points", data["tables"]["problem"])
    make_bar_chart(charts["ai"], "Trust in AI Weekly Plan", data["tables"]["ai_trust"])
    make_bar_chart(charts["features"], "Most Desired App Features", data["tables"]["useful_features"][:6], "percent_of_respondents")
    make_bar_chart(charts["barriers"], "Consistency Barriers from Open Answers", data["tables"]["hardest_themes"])

    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("GymBuddy User Research Analysis")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 122, 77)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Deliverable 1: Market & User Research | Sample size: 20 Google Form responses").italic = True

    add_callout(
        doc,
        "Executive Summary",
        "The research supports GymBuddy's core thesis: beginner and budget gym users need practical in-gym clarity more than advanced fitness optimization. The MVP should help users know what to do today, show how to do it safely, offer alternatives when equipment is busy, and adapt next week's plan from check-in behavior.",
        PALE_GREEN,
    )

    add_heading(doc, "1. Research Objective", 1)
    doc.add_paragraph(
        "The goal of this research was to understand the pain points of beginner and budget gym users, especially the moments that cause confusion, skipped exercises, and early drop-off. The findings are intended to guide the GymBuddy MVP scope, onboarding, workout flow, demo support, check-ins, and adaptive weekly planning."
    )

    add_heading(doc, "2. Methodology", 1)
    methodology_rows = [
        ["Method", "Google Form survey"],
        ["Responses", str(data["total_responses"])],
        ["Audience", "Current gym users, recently joined users, planning users, and users who stopped going"],
        ["Research focus", "Gym clarity, skipped exercises, equipment constraints, AI trust, desired app features, and consistency barriers"],
    ]
    add_table(doc, ["Field", "Details"], methodology_rows, [1.8, 4.7])

    add_heading(doc, "3. Key Metrics", 1)
    kpi_rows = [[k["metric"], k["display"]] for k in data["kpis"]]
    add_table(doc, ["Metric", "Result"], kpi_rows, [3.3, 3.2])

    add_heading(doc, "4. What Users Struggle With Inside the Gym", 1)
    doc.add_paragraph(
        "The biggest reported pain point is equipment availability, followed closely by uncertainty around correct form. This matters because the GymBuddy flow should not assume the user's first exercise option will always be available or that the user knows how to perform it correctly."
    )
    doc.add_picture(str(charts["pain"]), width=Inches(6.4))
    doc.add_paragraph("Figure 1. Biggest in-gym pain points from survey responses.").italic = True
    pain_rows = [[r["label"], r["count"], pct(r["percent"])] for r in data["tables"]["problem"]]
    add_table(doc, ["Pain Point", "Count", "% of Respondents"], pain_rows, [3.4, 1.2, 1.9])

    add_heading(doc, "5. AI Trust and Safety", 1)
    doc.add_paragraph(
        "Trust in AI is high, but conditional. Most users are willing to trust an AI weekly gym plan when it is simple, supported by demo videos, or includes safety warnings. This means AI should be positioned as a guided planning and adaptation layer, not as an unrestricted trainer replacement."
    )
    doc.add_picture(str(charts["ai"]), width=Inches(6.4))
    doc.add_paragraph("Figure 2. Trust in an AI-generated weekly plan.").italic = True

    add_heading(doc, "6. Desired Product Features", 1)
    doc.add_paragraph(
        "The most requested features are simple daily workouts, sets/reps/weight guidance, demo videos, and progress tracking. This validates the MVP direction: a simple workout checklist with clear instructions and low-friction logging."
    )
    doc.add_picture(str(charts["features"]), width=Inches(6.4))
    doc.add_paragraph("Figure 3. Most desired GymBuddy app features.").italic = True
    feature_rows = [
        [r["label"], r["count"], pct(r["percent_of_respondents"])]
        for r in data["tables"]["useful_features"][:6]
    ]
    add_table(doc, ["Feature", "Count", "% of Respondents"], feature_rows, [3.4, 1.2, 1.9])

    add_heading(doc, "7. Consistency Barriers", 1)
    doc.add_paragraph(
        "Open-ended responses show that consistency is affected by motivation, energy after work, difficulty starting, soreness, and practical friction such as distance or weather. GymBuddy should reduce mental effort before and during the workout."
    )
    doc.add_picture(str(charts["barriers"]), width=Inches(6.4))
    doc.add_paragraph("Figure 4. Themes from open-ended consistency-barrier answers.").italic = True

    add_heading(doc, "8. Sharp User Insights", 1)
    insight_rows = [
        [str(i + 1), insight["title"], insight["evidence"], insight["product_decision"]]
        for i, insight in enumerate(data["insights"])
    ]
    add_table(doc, ["#", "Insight", "Evidence", "Product Decision"], insight_rows, [0.35, 1.75, 2.15, 2.25])

    add_heading(doc, "9. Representative User Quotes", 1)
    selected_quotes = [
        q
        for q in data["quotes"]
        if q and q.lower().strip() not in {"motivation", "not any"}
    ][:8]
    for quote in selected_quotes:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(f'"{sanitize(quote)}"')
        r.italic = True
        r.font.color.rgb = RGBColor(36, 65, 91)

    add_heading(doc, "10. Implications for GymBuddy MVP", 1)
    decisions = [
        "Prioritize a mobile-first Today's Workout screen over a broad dashboard.",
        "Show exact exercise, sets, reps, rest timer, and weight logging for each movement.",
        "Include one short demo link and one equipment-busy alternative for every exercise.",
        "Keep check-ins under 20 seconds: completed, skipped, pain, energy, confidence.",
        "Use AI for weekly plan generation and next-week adaptation, with visible safety boundaries.",
        "Track Week 4 completion as the main success metric, supported by first check-in and confidence change as early indicators.",
    ]
    for decision in decisions:
        doc.add_paragraph(decision, style="List Bullet")

    add_callout(
        doc,
        "Recommended MVP Positioning",
        "GymBuddy should be positioned as a first-month gym confidence system for budget gym beginners. It should not try to be a complete fitness app. Its job is to help the user enter the gym, complete the next right workout, check in quickly, and come back next week.",
        PALE_ORANGE,
    )

    add_heading(doc, "Appendix: Source Summary", 1)
    source_rows = [
        ["Source file", data["source_file"]],
        ["Analysis date", "2026-06-22"],
        ["Prepared for", "GymBuddy Product Management Launchpad assignment"],
    ]
    add_table(doc, ["Item", "Value"], source_rows, [1.7, 4.8])

    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
