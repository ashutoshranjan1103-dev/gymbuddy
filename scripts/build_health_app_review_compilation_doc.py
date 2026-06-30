from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"C:\Users\ashut\OneDrive\Documents\Gym Buddy")
OUT_DIR = PROJECT / "outputs" / "secondary_research"
DOCX_PATH = OUT_DIR / "Top_Health_App_Store_Review_Compilation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "EAF5EF"
BORDER = "C9D3DF"


SOURCES = [
    ("S1", "Google Play Health & Fitness chart/category", "https://play.google.com/store/apps/category/HEALTH_AND_FITNESS?hl=en_US"),
    ("S2", "Apple App Store Health & Fitness chart", "https://apps.apple.com/us/iphone/charts/6013"),
    ("S3", "Healthify Google Play listing", "https://play.google.com/store/apps/details?id=com.healthifyme.basic"),
    ("S4", "Healthify App Store India listing", "https://apps.apple.com/in/app/healthify-ai-calorie-tracker/id943712366"),
    ("S5", "Healthify user feedback signals", "https://www.trustpilot.com/review/www.healthifyme.com"),
    ("S6", "cult.fit App Store India listing", "https://apps.apple.com/in/app/cult-fit-gym-workout-fitness/id1217794588"),
    ("S7", "cult.fit Google Play listing", "https://play.google.com/store/apps/details?id=fit.cure.android"),
    ("S8", "MyFitnessPal Google Play listing", "https://play.google.com/store/apps/details?id=com.myfitnesspal.android"),
    ("S9", "MyFitnessPal App Store reviews", "https://apps.apple.com/eg/app/myfitnesspal-calorie-counter/id341232718?see-all=reviews"),
    ("S10", "MyFitnessPal official testimonials", "https://www.myfitnesspal.com/"),
    ("S11", "Strava App Store listing", "https://apps.apple.com/us/app/strava-run-bike-walk/id426826309"),
    ("S12", "Strava Google Play listing", "https://play.google.com/store/apps/details?id=com.strava"),
    ("S13", "Google Health/Fitbit Google Play listing", "https://play.google.com/store/apps/details?id=com.fitbit.FitbitMobile"),
    ("S14", "Google Health/Fitbit App Store listing", "https://apps.apple.com/us/app/google-health-fitbit/id462638897"),
    ("S15", "Google Health user reaction coverage", "https://www.techradar.com/health-fitness/fitness-apps/google-health-is-getting-heat-for-being-unbelievably-bad-after-replacing-the-fitbit-app-but-google-says-fixes-are-coming"),
    ("S16", "Samsung Health Google Play listing", "https://play.google.com/store/apps/details?id=com.sec.android.app.shealth"),
    ("S17", "Samsung Health App Store reviews", "https://apps.apple.com/ca/app/samsung-health/id1224541484?platform=iphone&see-all=reviews"),
    ("S18", "Flo App Store listing", "https://apps.apple.com/us/app/flo-cycle-period-tracker/id1038369065"),
    ("S19", "Flo Google Play listing", "https://play.google.com/store/apps/details?id=org.iggymedia.periodtracker"),
    ("S20", "Calm Google Play listing", "https://play.google.com/store/apps/details?id=com.calm.android"),
    ("S21", "Calm App Store reviews", "https://apps.apple.com/us/app/calm/id571800810?see-all=reviews"),
    ("S22", "Calm official review page", "https://www.calm.com/"),
    ("S23", "BetterMe Google Play listing", "https://play.google.com/store/apps/details?id=com.gen.workoutme"),
    ("S24", "BetterMe App Store reviews", "https://apps.apple.com/kz/app/betterme-health-coaching/id1264546236"),
    ("S25", "Nike Training Club Google Play listing", "https://play.google.com/store/apps/details?id=com.nike.ntc"),
    ("S26", "Nike Training Club App Store listing", "https://apps.apple.com/us/app/nike-training-club/id301521403"),
]


APPS = [
    {
        "app": "Healthify",
        "market": "India + global",
        "positive": "Users like India-friendly calorie tracking, AI diet guidance, coach/dietitian support, weight-loss accountability, and a large food database.",
        "negative": "Common complaints cluster around premium upsell, coach/support continuity, glitches, and a UI that can feel pushy or cluttered.",
        "lesson": "For GymBuddy: keep onboarding and logging light; do not bury the basic habit loop behind payment friction.",
        "sources": "[S3][S4][S5]",
    },
    {
        "app": "cult.fit",
        "market": "India",
        "positive": "Users value the one-stop fitness ecosystem: gyms, group workouts, yoga, dance, home fitness, and local brand trust.",
        "negative": "Review signals point to device/feature issues, occasional experience inconsistency, and a broad app that may not guide solo budget-gym training deeply.",
        "lesson": "GymBuddy can win by being narrower: daily gym-floor clarity for beginners, not a full fitness marketplace.",
        "sources": "[S6][S7]",
    },
    {
        "app": "MyFitnessPal",
        "market": "Global",
        "positive": "Users praise calorie/macro tracking, food database breadth, weight-loss accountability, integrations, and habit awareness.",
        "negative": "Complaints include unverified duplicate food entries, logging friction, crashes during food entry, ads, and features moving behind subscription.",
        "lesson": "Tracking works when it is fast and trustworthy. GymBuddy should avoid making workout logging feel like homework.",
        "sources": "[S8][S9][S10]",
    },
    {
        "app": "Strava",
        "market": "Global",
        "positive": "Users like GPS tracking, community motivation, activity sharing, segment competition, challenges, and multi-sport support.",
        "negative": "Complaints often center on subscription value, paywalled features, privacy concerns, and GPS/sync issues.",
        "lesson": "Community can motivate, but GymBuddy's first version should focus on confidence and completion before social features.",
        "sources": "[S11][S12]",
    },
    {
        "app": "Google Health / Fitbit",
        "market": "Global + India rollout",
        "positive": "Users value wearable-linked health metrics, steps, sleep, activity trends, cardiac-rehab style progress tracking, and customizable dashboards.",
        "negative": "Recent rebrand feedback highlights syncing problems, missing/changed features, inaccurate data display, clunky layout, and unwanted AI prominence.",
        "lesson": "Data trust is fragile. GymBuddy should make adaptation explainable and never hide core progress behind confusing redesigns.",
        "sources": "[S13][S14][S15]",
    },
    {
        "app": "Samsung Health",
        "market": "Global, strong Android/Samsung base",
        "positive": "Users like automatic activity tracking, steps, sleep, calories, exercise records, and Samsung wearable integration.",
        "negative": "Negative review themes include sync failures, removed/changed features, watch compatibility issues, and occasional sleep/step accuracy complaints.",
        "lesson": "GymBuddy should not depend on wearables in the MVP; manual check-ins are simpler and more reliable for budget users.",
        "sources": "[S16][S17]",
    },
    {
        "app": "Flo",
        "market": "Global women's health",
        "positive": "Users value cycle prediction, pregnancy tracking, reminders, symptom logging, and doctor-backed educational content.",
        "negative": "Complaints cluster around paywalls, subscription pressure, limited free pregnancy/cycle content, and privacy concerns around sensitive health data.",
        "lesson": "Sensitive health data requires trust. GymBuddy should ask only necessary questions and explain how data improves the plan.",
        "sources": "[S18][S19]",
    },
    {
        "app": "Calm",
        "market": "Global mental wellness",
        "positive": "Users praise sleep stories, meditation programs, anxiety support, relaxation, and therapist-adjacent support between sessions.",
        "negative": "Common issues include subscription/cancellation friction, customer service complaints, cluttered interface, and content refresh expectations.",
        "lesson": "Motivation features should feel supportive, not noisy. GymBuddy can use calming recovery prompts without overloading the user.",
        "sources": "[S20][S21][S22]",
    },
    {
        "app": "BetterMe",
        "market": "Global",
        "positive": "Users like beginner-friendly guided workouts, simple routines, personalization, meal ideas, and motivation to restart exercise.",
        "negative": "Complaints focus strongly on cost, trial conversion, auto-renewal, cancellation/refund issues, and billing trust.",
        "lesson": "For budget gym users, pricing trust matters. A student MVP should keep the promise free/simple while testing retention.",
        "sources": "[S23][S24]",
    },
    {
        "app": "Nike Training Club",
        "market": "Global",
        "positive": "Users like free trainer-led workouts, strong brand trust, beginner-to-advanced programs, yoga, strength, recovery, and polished content.",
        "negative": "Negative signals include app bugs, missing workout tabs/categories, crashes, search/navigation friction, and limited gym-floor logging.",
        "lesson": "Great content is not enough. GymBuddy should make today's workout easier to execute, log, and adapt in a real gym.",
        "sources": "[S25][S26]",
    },
]


THEMES = [
    ["Easy tracking and progress visibility", "Users repeatedly like apps that make steps, food, workouts, sleep, or cycles visible.", "Show completed days, confidence, and next workout clearly."],
    ["Guidance reduces anxiety", "Positive reviews praise coach-like guidance, programs, demos, and reminders.", "Give beginner-friendly plan logic and demo links."],
    ["Subscription friction creates distrust", "Many negative reviews complain when core value feels paywalled, auto-renewed, or hard to cancel.", "Keep the MVP transparent and free during testing."],
    ["Bugs destroy habit loops", "Sync errors, crashes, missing tabs, inaccurate data, and bad redesigns appear across categories.", "Workout check-in must work reliably in a mobile browser."],
    ["Too much complexity hurts beginners", "Large libraries and broad ecosystems can overwhelm users who only need the next action.", "Prioritize today's workout over dashboards and large menus."],
]


def sanitize(text):
    text = "" if text is None else str(text)
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("ascii", "ignore").decode("ascii").strip()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, size=8.1, bold=False, color=INK):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.04
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(color)
            if bold:
                run.font.bold = True


def add_table(doc, headers, rows, widths, font_size=8.1):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = sanitize(header)
        set_cell_shading(cell, LIGHT_BLUE)
        style_cell_text(cell, size=8.3, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = sanitize(value)
            style_cell_text(cells[idx], size=font_size)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = sanitize(text)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_reference_table(doc):
    doc.add_heading("Sources", level=1)
    doc.add_paragraph("Source IDs are used in the review table. App store rankings and reviews change frequently, so this should be treated as a current public-review snapshot.")
    table = doc.add_table(rows=1, cols=4)
    widths = [0.45, 2.0, 3.0, 1.05]
    set_table_width(table, widths)
    for idx, header in enumerate(["ID", "Source", "URL Type", "Link"]):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        style_cell_text(cell, size=8, bold=True)
    for sid, title, url in SOURCES:
        cells = table.add_row().cells
        cells[0].text = sid
        cells[1].text = sanitize(title)
        cells[2].text = "App store / public review source"
        add_hyperlink(cells[3].paragraphs[0], "Open", url)
        for cell in cells[:3]:
            style_cell_text(cell, size=7.5)
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table, widths)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("GymBuddy - App Review Scan")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Common Positive & Negative App Store Reviews")
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Top health and fitness apps used in India and globally | Review snapshot for GymBuddy")
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor.from_string(MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_width(callout, [6.5])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREEN)
    cell.text = "Selection method: Apps were selected from current public Health & Fitness chart/category signals plus widely used India/global health apps relevant to GymBuddy's market. This is not a paid Sensor Tower-style ranking; it is a practical PM review scan using public app-store and review sources. [S1][S2]"
    style_cell_text(cell, size=9.2)
    doc.add_paragraph()

    doc.add_heading("1. Review Matrix", level=1)
    rows = [[a["app"], a["market"], a["positive"], a["negative"], a["lesson"], a["sources"]] for a in APPS]
    add_table(
        doc,
        ["App", "Market", "Common Positive Reviews", "Common Negative Reviews", "PM Lesson For GymBuddy", "Sources"],
        rows,
        [0.85, 0.75, 1.55, 1.55, 1.45, 0.35],
        font_size=7.1,
    )

    doc.add_heading("2. Cross-App Review Themes", level=1)
    add_table(doc, ["Theme", "What Users Say Across Apps", "Implication For GymBuddy"], THEMES, [1.6, 2.55, 2.35], font_size=8.3)

    doc.add_heading("3. One-Line Conclusion", level=1)
    doc.add_paragraph(
        "The biggest review pattern is simple: users reward apps that make progress visible and action easy, but punish apps that add paywall friction, bugs, confusing redesigns, or heavy logging. GymBuddy should therefore win on beginner clarity, reliable mobile-browser check-ins, simple workout logging, and transparent AI adaptation."
    )

    add_reference_table(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
