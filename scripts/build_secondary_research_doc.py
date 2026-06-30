from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT = Path(r"C:\Users\ashut\OneDrive\Documents\Gym Buddy")
OUT_DIR = PROJECT / "outputs" / "secondary_research"
CHART_DIR = OUT_DIR / "charts"
DOCX_PATH = OUT_DIR / "GymBuddy_Secondary_Research_Report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "EAF5EF"
LIGHT_YELLOW = "FFF7E6"
LIGHT_RED = "FDECEC"
BORDER = "C9D3DF"
WHITE = "FFFFFF"


SOURCES = [
    {
        "id": "S1",
        "title": "mHealth Apps Market Size, Share & Trends Analysis Report",
        "publisher": "Grand View Research",
        "url": "https://www.grandviewresearch.com/industry-analysis/mhealth-app-market",
        "note": "Reports global mHealth apps market size and projected CAGR.",
    },
    {
        "id": "S2",
        "title": "Fitness Apps Market Size, Share & Trends Analysis Report",
        "publisher": "Grand View Research",
        "url": "https://www.grandviewresearch.com/industry-analysis/fitness-app-market",
        "note": "Reports global fitness app market size and projected CAGR.",
    },
    {
        "id": "S3",
        "title": "India Fitness Market Report 2025",
        "publisher": "Health & Fitness Association and Deloitte India",
        "url": "https://www.healthandfitness.org/wp-content/uploads/India-Fitness-Market-2025_FINAL_Print.pdf",
        "note": "Reports commercial fitness facilities, sector growth, and India fitness market context.",
    },
    {
        "id": "S4",
        "title": "India's fitness market to double by 2030",
        "publisher": "Deloitte India",
        "url": "https://www.deloitte.com/in/en/about/press-room/indias-fitness-market-to-double-by-2030-per-a-deloitte-and-hfa-report.html",
        "note": "Reports expected membership growth from 2024 to 2030 and low fitness penetration.",
    },
    {
        "id": "S5",
        "title": "Nearly 1.8 billion adults at risk of disease from not doing enough physical activity",
        "publisher": "World Health Organization",
        "url": "https://www.who.int/news/item/26-06-2024-nearly-1.8-billion-adults-at-risk-of-disease-from-not-doing-enough-physical-activity",
        "note": "Reports global physical inactivity levels and trend.",
    },
    {
        "id": "S6",
        "title": "Global Strategy on Digital Health 2020-2025",
        "publisher": "World Health Organization",
        "url": "https://www.who.int/publications/i/item/9789240020924",
        "note": "Frames digital health as a way to strengthen health systems when designed and implemented responsibly.",
    },
    {
        "id": "S7",
        "title": "Can Smartphone Apps Increase Physical Activity? Systematic Review and Meta-Analysis",
        "publisher": "JMIR mHealth and uHealth",
        "url": "https://www.jmir.org/2019/3/e12053/",
        "note": "Finds modest evidence that smartphone apps can increase physical activity, especially short term.",
    },
    {
        "id": "S8",
        "title": "Effectiveness of physical activity monitors in adults: systematic review and meta-analysis",
        "publisher": "British Journal of Sports Medicine",
        "url": "https://bjsm.bmj.com/content/55/8/422",
        "note": "Reports positive effects of smartphone apps and activity trackers on physical activity.",
    },
    {
        "id": "S9",
        "title": "Effect of smartphone applications on physical activity in healthy young adults",
        "publisher": "PLOS ONE",
        "url": "https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0301088",
        "note": "Finds no significant pooled effect in healthy young adults and cautions that the evidence base is limited.",
    },
    {
        "id": "S10",
        "title": "Nike Training Club App Store listing",
        "publisher": "Apple App Store",
        "url": "https://apps.apple.com/us/app/nike-training-club/id301521403",
        "note": "Official iOS store listing and public user review snippets.",
    },
    {
        "id": "S11",
        "title": "Nike Training Club official product page",
        "publisher": "Nike",
        "url": "https://www.nike.com/ntc-app",
        "note": "Official description of free workouts, trainers, programs, and guidance.",
    },
    {
        "id": "S12",
        "title": "cult.fit App Store listing",
        "publisher": "Apple App Store India",
        "url": "https://apps.apple.com/in/app/cult-fit-gym-workout-fitness/id1217794588",
        "note": "Official iOS store listing and public rating/review signals.",
    },
    {
        "id": "S13",
        "title": "cult.fit Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=fit.cure.android",
        "note": "Official Android listing and public review snippets.",
    },
    {
        "id": "S14",
        "title": "Healthify Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.healthifyme.basic",
        "note": "Official Android listing for AI diet, health, and fitness features.",
    },
    {
        "id": "S15",
        "title": "Healthify official app page",
        "publisher": "Healthify",
        "url": "https://www.healthifyme.com/",
        "note": "Company claims and feature summary for AI coaching, food tracking, and health plans.",
    },
    {
        "id": "S16",
        "title": "MyFitnessPal App Store listing",
        "publisher": "Apple App Store",
        "url": "https://apps.apple.com/us/app/myfitnesspal-calorie-counter/id341232718",
        "note": "Official iOS listing for calorie, nutrition, and meal tracking.",
    },
    {
        "id": "S17",
        "title": "MyFitnessPal Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.myfitnesspal.android",
        "note": "Official Android listing for nutrition goals and tracking features.",
    },
    {
        "id": "S18",
        "title": "Fitbod App Store listing",
        "publisher": "Apple App Store",
        "url": "https://apps.apple.com/us/app/fitbod-gym-fitness-planner/id1041517543",
        "note": "Official iOS listing for AI workout planning and guided programs.",
    },
    {
        "id": "S19",
        "title": "Fitbod Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.fitbod.fitbod",
        "note": "Official Android listing for personalized workouts by goal, equipment, and fitness level.",
    },
    {
        "id": "S20",
        "title": "Strong official product page",
        "publisher": "Strong",
        "url": "https://www.strong.app/",
        "note": "Official workout tracker positioning, user scale, and ratings claims.",
    },
    {
        "id": "S21",
        "title": "Strong App Store listing",
        "publisher": "Apple App Store",
        "url": "https://apps.apple.com/us/app/strong-workout-tracker-gym-log/id464254577",
        "note": "Official iOS listing and public review signals for workout logging.",
    },
    {
        "id": "S22",
        "title": "JEFIT App Store listing",
        "publisher": "Apple App Store",
        "url": "https://apps.apple.com/us/app/jefit-workout-plan-gym-tracker/id449810000",
        "note": "Official iOS listing for workout planning, exercise library, and tracking.",
    },
    {
        "id": "S23",
        "title": "JEFIT Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=je.fit",
        "note": "Official Android listing and public review signals.",
    },
    {
        "id": "S24",
        "title": "Freeletics Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.freeletics.lite",
        "note": "Official Android listing and public review snippets for coaching and workout planning.",
    },
    {
        "id": "S25",
        "title": "Caliber Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.caliberfitness.app",
        "note": "Official Android listing and review snippets for strength training planning and tracking.",
    },
    {
        "id": "S26",
        "title": "Hevy Google Play listing",
        "publisher": "Google Play",
        "url": "https://play.google.com/store/apps/details?id=com.hevy",
        "note": "Official Android listing and review snippets for gym logging, graphs, and social motivation.",
    },
]


def sanitize(text):
    text = "" if text is None else str(text)
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
        "\u20b9": "Rs ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("ascii", "ignore").decode("ascii").strip()


def image_font(size=28, bold=False):
    candidates = [
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap(draw, text, max_width, font):
    words = sanitize(text).split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def make_market_chart(path):
    rows = [
        ("Global mHealth apps market CAGR", 14.8, "[S1]"),
        ("Global fitness apps market CAGR", 13.4, "[S2]"),
        ("India fitness market CAGR to 2030", 15.0, "[S3]"),
        ("India gym membership growth 2024-2030", 88.6, "[S4]"),
    ]
    width, height = 1300, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = image_font(36, True)
    label_font = image_font(24, False)
    value_font = image_font(24, True)
    small_font = image_font(20, False)
    draw.text((50, 35), "Market Signals Supporting GymBuddy", fill=f"#{INK}", font=title_font)
    draw.text(
        (50, 82),
        "Growth signals are strongest where digital health, fitness apps, and India's gym market overlap.",
        fill=f"#{MUTED}",
        font=small_font,
    )
    left_label = 55
    left_bar = 520
    top = 150
    max_val = 90
    bar_max = 650
    colors = [BLUE, DARK_BLUE, "1E7A52", "B76E00"]
    for idx, (label, value, source) in enumerate(rows):
        y = top + idx * 115
        for line_no, line in enumerate(wrap(draw, label, 405, label_font)[:2]):
            draw.text((left_label, y + line_no * 27), line, fill=f"#{INK}", font=label_font)
        draw.rounded_rectangle((left_bar, y, left_bar + bar_max, y + 44), radius=10, fill="#EEF2F6")
        bar_len = int((value / max_val) * bar_max)
        draw.rounded_rectangle((left_bar, y, left_bar + bar_len, y + 44), radius=10, fill=f"#{colors[idx]}")
        suffix = "growth" if "membership" in label.lower() else "CAGR"
        draw.text((left_bar + bar_len + 18, y + 7), f"{value:.1f}% {suffix} {source}", fill=f"#{INK}", font=value_font)
    draw.text((50, height - 45), "Note: Membership growth calculated from 12.3M to 23.2M reported by Deloitte India.", fill=f"#{MUTED}", font=small_font)
    img.save(path)


def make_evidence_chart(path):
    rows = [
        ("Need: physical inactivity remains high", 5, "WHO reports nearly 1.8B inactive adults [S5]"),
        ("Opportunity: mobile tools can help", 4, "Meta-analyses show positive/modest activity effects [S7][S8]"),
        ("Caution: impact is not automatic", 3, "Some young-adult app trials show no significant effect [S9]"),
        ("PM lesson: retention is the product", 5, "Apps need guidance, feedback, and low friction, not only content."),
    ]
    width, height = 1300, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = image_font(36, True)
    label_font = image_font(24, True)
    body_font = image_font(21, False)
    draw.text((50, 35), "What The Evidence Means For Product Design", fill=f"#{INK}", font=title_font)
    y = 120
    for idx, (label, score, note) in enumerate(rows):
        card_y = y + idx * 135
        fill = ["#F5FAFF", "#EFF8F3", "#FFF7E8", "#F4F6F9"][idx]
        draw.rounded_rectangle((50, card_y, 1250, card_y + 105), radius=16, fill=fill, outline="#D9E3EA")
        draw.text((78, card_y + 18), label, fill=f"#{INK}", font=label_font)
        draw.text((78, card_y + 55), sanitize(note), fill=f"#{MUTED}", font=body_font)
        x0 = 1040
        for dot in range(5):
            color = "#2E74B5" if dot < score else "#D5DCE5"
            draw.ellipse((x0 + dot * 34, card_y + 38, x0 + dot * 34 + 20, card_y + 58), fill=color)
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_paragraph(paragraph, size=10, color=INK, bold=False):
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = bold if bold else run.font.bold


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = sanitize(header)
        set_cell_shading(header_cells[idx], header_fill)
        for p in header_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor.from_string(INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = sanitize(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"
                    run.font.color.rgb = RGBColor.from_string(INK)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(sanitize(text), level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
    return p


def add_callout(doc, title, body, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    title_run = p.add_run(sanitize(title))
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run("\n" + sanitize(body))
    for run in p.runs:
        run.font.name = "Calibri"
        if run.font.size is None:
            run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(sanitize(item))


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = sanitize(text)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("GymBuddy Secondary Research")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_reference_list(doc):
    add_heading(doc, "Appendix: References", 1)
    doc.add_paragraph(
        "The source IDs below are used throughout the document. Store ratings and review snippets are volatile; they should be treated as public signals observed during the research scan, not permanent values."
    )
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [0.45, 1.35, 3.75, 0.95])
    headers = ["ID", "Publisher", "Source", "Link"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = sanitize(header)
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(INK)

    for source in SOURCES:
        cells = table.add_row().cells
        for idx, value in enumerate([source["id"], source["publisher"], source["title"]]):
            cells[idx].text = sanitize(value)
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7.4)
                    run.font.color.rgb = RGBColor.from_string(INK)
        link_p = cells[3].paragraphs[0]
        link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(link_p, "Open", source["url"])

    set_table_width(table, [0.45, 1.35, 3.75, 0.95])
    doc.add_paragraph()


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    market_chart = CHART_DIR / "market_signals.png"
    evidence_chart = CHART_DIR / "evidence_implications.png"
    make_market_chart(market_chart)
    make_evidence_chart(evidence_chart)

    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("GymBuddy Secondary Research Report")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Market scope, health-impact evidence, competitive app review scan, and product implications")
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle_run.font.size = Pt(11)
    doc.add_paragraph("Prepared for: GymBuddy Product Management Capstone | Date: 22 June 2026")

    add_callout(
        doc,
        "Executive Summary",
        "GymBuddy has a credible opportunity because the fitness market is growing, physical inactivity remains a public-health problem, and existing apps often optimize for content, tracking, or advanced users rather than the first-month budget gym beginner. The evidence suggests that health apps can improve behavior, but impact is not automatic; the product must reduce gym-floor confusion, make completion easy, and adapt weekly based on real check-ins.",
        LIGHT_GREEN,
    )

    add_heading(doc, "1. Research Questions", 1)
    add_bullets(
        doc,
        [
            "Is there a meaningful business opportunity for gym and health apps, especially in India?",
            "Do gym apps and health apps actually improve health behavior in real life?",
            "What do popular fitness apps do well, and where do users still feel friction?",
            "What product gaps should GymBuddy exploit for beginner budget gym users?",
        ],
    )

    add_heading(doc, "2. Market And Business Scope", 1)
    doc.add_paragraph(
        "The business case for GymBuddy sits at the intersection of three trends: growth in global mHealth apps, growth in fitness-specific apps, and the expansion of India's commercial fitness market."
    )
    doc.add_picture(str(market_chart), width=Inches(6.4))
    p = doc.add_paragraph("Figure 1. Market signals that support the GymBuddy opportunity. Sources: [S1][S2][S3][S4].")
    p.italic = True

    market_rows = [
        [
            "Global mHealth apps",
            "Grand View Research estimates the global mHealth apps market at USD 37.5B in 2024 and projects USD 86.37B by 2030 at 14.8% CAGR.",
            "Health behavior is becoming a software category, not only a clinic or trainer category.",
            "[S1]",
        ],
        [
            "Global fitness apps",
            "Grand View Research estimates the global fitness apps market at USD 12.1B in 2025 and projects USD 33.6B by 2033 at 13.4% CAGR.",
            "Fitness app demand is large enough for focused, use-case-specific products.",
            "[S2]",
        ],
        [
            "India commercial fitness",
            "HFA/Deloitte report about 46,500 fitness facilities in India and a fitness market expected to grow about 15% CAGR to 2030.",
            "India is not under-supplied by gyms; the gap is guidance after users enter the gym.",
            "[S3]",
        ],
        [
            "India gym membership",
            "Deloitte India reports memberships expected to rise from 12.3M in 2024 to 23.2M by 2030, while penetration remains low.",
            "A beginner-first product can ride gym adoption while solving early drop-off.",
            "[S4]",
        ],
    ]
    add_table(doc, ["Area", "Evidence", "Business Meaning", "Source"], market_rows, [1.2, 2.55, 2.15, 0.6], font_size=8.2)

    add_callout(
        doc,
        "Business Perspective",
        "The strongest wedge is not 'another workout library'. It is a guided first-month retention layer for budget gyms: onboarding, daily workout clarity, demo support, check-ins, and next-week adaptation. This can later become a B2C app, a gym-partner retention tool, or a low-cost AI coaching layer for budget facilities.",
        LIGHT_YELLOW,
    )

    add_heading(doc, "3. Do Health Apps Impact Real Life Health?", 1)
    doc.add_paragraph(
        "The evidence is positive but nuanced. Digital tools can improve activity when they create reminders, feedback loops, goals, and self-monitoring. However, simply giving users content is not enough, and the effect can weaken when apps are too generic or require too much effort."
    )
    doc.add_picture(str(evidence_chart), width=Inches(6.4))
    p = doc.add_paragraph("Figure 2. Evidence implications for GymBuddy's product design. Sources: [S5][S6][S7][S8][S9].")
    p.italic = True

    health_rows = [
        [
            "Public-health need",
            "WHO reported that nearly one-third of adults, about 1.8B people, did not meet recommended physical-activity levels in 2022.",
            "There is a real behavior problem, not just a content-discovery problem.",
            "[S5]",
        ],
        [
            "Digital-health legitimacy",
            "WHO's digital-health strategy positions digital health as a tool for strengthening health systems when implemented responsibly.",
            "GymBuddy should present AI as supportive guidance with safe boundaries.",
            "[S6]",
        ],
        [
            "Positive app evidence",
            "A JMIR systematic review/meta-analysis found modest evidence that smartphone apps can increase physical activity, especially over short periods.",
            "A 4-week beginner journey is a sensible MVP horizon.",
            "[S7]",
        ],
        [
            "Tracker/app effectiveness",
            "A BJSM meta-analysis found smartphone apps and activity trackers had positive effects on physical activity in adults.",
            "Self-monitoring and feedback should be core, not optional.",
            "[S8]",
        ],
        [
            "Important caution",
            "A 2024 PLOS ONE review of healthy young adults found no significant pooled effect and noted the limited number of studies.",
            "GymBuddy must prove impact through check-ins, completion, and user testing.",
            "[S9]",
        ],
    ]
    add_table(doc, ["Evidence Area", "Finding", "Implication For GymBuddy", "Source"], health_rows, [1.25, 2.45, 2.25, 0.55], font_size=8.2)

    add_heading(doc, "4. Competitive App Store Review Scan", 1)
    doc.add_paragraph(
        "This scan uses official app store pages, official product pages, and public review snippets visible during the research pass. App ratings and individual reviews change frequently, so the conclusions focus on stable product patterns rather than exact rating values."
    )

    competitive_rows = [
        [
            "Nike Training Club",
            "Free workout library; trainer-led strength, conditioning, yoga, pilates, recovery, and programs.",
            "Excellent free content and brand trust. Good for users who want guided classes without paying.",
            "Less focused on the budget-gym floor. Review snippets mention app crashes/navigation issues and metrics visibility complaints.",
            "GymBuddy can be more practical inside a real gym: exact sets, equipment alternatives, and check-ins.",
            "[S10][S11]",
        ],
        [
            "cult.fit",
            "India-focused fitness ecosystem across gyms, group classes, home workouts, yoga, dance, and services.",
            "Strong local brand, broad supply network, and high store rating signals.",
            "The app is broad. Review snippets mention feature/device issues, and beginners may still need exercise-level guidance inside a budget gym.",
            "GymBuddy should stay narrower: first-month confidence for people already paying for a budget gym.",
            "[S12][S13]",
        ],
        [
            "Healthify",
            "AI diet, health tracking, calorie/nutrition guidance, coaching, and Indian food database positioning.",
            "Strong nutrition and weight-management proposition; India-relevant food context.",
            "More diet/weight-loss oriented than gym execution. Gym beginners still need form, sets, reps, and equipment swaps.",
            "GymBuddy can integrate simple diet tips later, but the core wedge should be workout completion.",
            "[S14][S15]",
        ],
        [
            "MyFitnessPal",
            "Calorie, macro, meal, water, and goal tracking.",
            "Very strong nutrition database and habit tracking. Useful for users already motivated to log food.",
            "Food logging can feel heavy; not a gym-floor workout coach. Some reviews complain about ads/paywall or free-feature limits.",
            "GymBuddy should keep logging lightweight and avoid asking beginners to track everything.",
            "[S16][S17]",
        ],
        [
            "Fitbod",
            "AI-personalized strength workouts based on goals, equipment, recovery, and progressive overload.",
            "Closest competitor for adaptive strength planning. Strong planning logic and equipment personalization.",
            "Can feel optimized for users who already understand exercises and gym structure; subscription/value concerns may appear in reviews.",
            "GymBuddy should beat it on beginner language, budget-gym assumptions, and first-month coaching.",
            "[S18][S19]",
        ],
        [
            "Strong",
            "Workout tracking, history, timer, templates, and gym logbook.",
            "Excellent for lifters who know what they plan to do and want clean logging.",
            "Not a beginner coach. It tracks the workout; it does not solve 'what should I do today and how do I do it?'",
            "GymBuddy should combine plan + demo + logging rather than logging alone.",
            "[S20][S21]",
        ],
        [
            "JEFIT",
            "Exercise database, premade plans, muscle breakdowns, workout tracking, and progress tracking.",
            "Large exercise library and plan ecosystem.",
            "Users can feel overwhelmed. Review snippets point to interface and sync/friction complaints.",
            "GymBuddy should avoid library overload and make the next action obvious.",
            "[S22][S23]",
        ],
        [
            "Freeletics",
            "AI coaching and workout planning, often bodyweight/home-workout oriented.",
            "Strong coaching identity and structured training.",
            "Review snippets mention schedule rigidity and limited flexibility around moving workouts.",
            "GymBuddy should adapt when users skip, miss days, or equipment is unavailable.",
            "[S24]",
        ],
        [
            "Caliber",
            "Strength training plans, coaching, workout history, and progress tracking.",
            "Good consistency/history value for strength users.",
            "Review snippets suggest it may be less useful above beginner plans and can miss exercises/equipment.",
            "GymBuddy should build equipment-aware alternatives from day one.",
            "[S25]",
        ],
        [
            "Hevy",
            "Workout logging, exercise videos, graphs, and social motivation.",
            "Strong gym logging and progress visualization.",
            "Social/logging value is stronger than first-month guided coaching.",
            "GymBuddy should make users feel safe and capable before asking them to become serious loggers.",
            "[S26]",
        ],
    ]
    add_table(
        doc,
        ["App", "Positioning", "Pros", "Cons / Review Signals", "Opportunity For GymBuddy", "Sources"],
        competitive_rows,
        [0.85, 1.3, 1.35, 1.45, 1.3, 0.55],
        font_size=7.2,
    )

    add_heading(doc, "5. Competitive Pattern Summary", 1)
    pattern_rows = [
        [
            "Content-heavy apps",
            "Nike Training Club, JEFIT",
            "They give many workouts/exercises, but beginners can still struggle with what to choose and how to execute in a crowded gym.",
        ],
        [
            "Ecosystem apps",
            "cult.fit",
            "They solve access and classes, but may not be optimized for a budget-gym user training alone.",
        ],
        [
            "Nutrition-first apps",
            "Healthify, MyFitnessPal",
            "They are useful for health goals but do not directly solve gym-floor confusion.",
        ],
        [
            "Tracker-first apps",
            "Strong, Hevy",
            "They work best after a user already has workout knowledge and a stable routine.",
        ],
        [
            "AI/adaptive strength apps",
            "Fitbod, Freeletics, Caliber",
            "They validate AI planning, but GymBuddy can differentiate through beginner language, Indian budget-gym assumptions, and low-friction weekly adaptation.",
        ],
    ]
    add_table(doc, ["Pattern", "Examples", "Meaning"], pattern_rows, [1.45, 1.55, 3.5], font_size=8.4)

    add_heading(doc, "6. Sharp Insights For GymBuddy", 1)
    insights = [
        (
            "The gym is already paid for; confidence is the missing layer.",
            "India's gym market is expanding, but the user's pain starts after entry. A budget gym user does not need another generic fitness feed; they need to know exactly what to do today.",
        ),
        (
            "Health apps can work when they change behavior loops, not when they only publish content.",
            "The scientific evidence supports apps that create feedback, self-monitoring, and action prompts. This validates GymBuddy's plan + demo + check-in + adaptation loop.",
        ),
        (
            "The best competitor gap is 'first-month gym floor execution'.",
            "Popular apps are strong in libraries, diet, classes, AI planning, or logging. GymBuddy should own beginner execution: alternatives, form links, timer, sets/reps/weight, and adaptive next week.",
        ),
    ]
    for title_text, body in insights:
        add_callout(doc, title_text, body, LIGHT_BLUE)

    add_heading(doc, "7. Recommended MVP Direction", 1)
    doc.add_paragraph(
        "For a top-scoring capstone, GymBuddy should be framed as a focused MVP rather than a large fitness platform. The product should prove that beginner budget gym users complete more workouts when the app reduces ambiguity at the exact moment of training."
    )
    add_bullets(
        doc,
        [
            "Core promise: 'Walk into the gym knowing exactly what to do today.'",
            "Primary user: Indian budget gym beginner in the first 30 days after joining.",
            "Primary flow: onboarding -> weekly AI plan -> today's workout -> demo link -> set/reps/weight/timer -> 20-second check-in -> next-week adaptation.",
            "Main metric: Week 4 completion, measured by users who check in at least once during Week 4.",
            "Differentiator: simple gym-floor confidence, not advanced analytics.",
        ],
    )

    mvp_rows = [
        [
            "Onboarding",
            "Ask goal, experience, days/week, height, weight, gender, gym type, equipment confidence, injuries, and preferred workout time.",
            "Collects only what is needed to generate a safe first-week plan.",
        ],
        [
            "Today's Workout",
            "Show warm-up, 5-7 exercises, sets/reps, rest timer, weight input, and one demo link per exercise.",
            "Reduces decision fatigue inside the gym.",
        ],
        [
            "Equipment Alternative",
            "Every exercise has a basic-gym and equipped-gym option.",
            "Solves crowded or limited gym environments.",
        ],
        [
            "Check-In",
            "Completed/skipped, difficulty, pain flag, confidence, and notes.",
            "Creates adaptation data without heavy logging.",
        ],
        [
            "Adaptive Week",
            "Next week changes volume, exercise difficulty, cardio time, or recovery based on completion and pain.",
            "Turns AI from a buzzword into a visible product behavior.",
        ],
    ]
    add_table(doc, ["MVP Area", "What To Build", "Why It Matters"], mvp_rows, [1.25, 3.15, 2.1], font_size=8.5)

    add_heading(doc, "8. How To Present This In Your Assignment", 1)
    add_bullets(
        doc,
        [
            "Show that the market is attractive, but avoid claiming market size alone proves product demand.",
            "Use health-impact research to justify behavior loops: plan, reminder, self-monitoring, feedback, adaptation.",
            "Use competitor analysis to show a clear wedge: beginner budget-gym execution, not generic workouts.",
            "Tie every feature to Week 4 completion, because that is the assignment's success metric.",
            "In user testing, measure whether users understand today's workout in under 60 seconds and complete the check-in in under 20 seconds.",
        ],
    )

    add_reference_list(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
